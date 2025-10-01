#!/usr/bin/env python3
"""
Create a sample .docx document for testing the AI Document Agent
"""

try:
    from docx import Document
    from pathlib import Path
    
    def create_sample_document():
        """Create a sample document with intentional errors for testing"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Sample Document for AI Processing', 0)
        
        # Add content with intentional errors
        content = [
            "This is a sample document for testing the AI Document Agent.",
            "",
            "The document contains some intentional errors and areas for improvement:",
            "",
            "• i think this system will work great",
            "• there are some grammar mistakes here", 
            "• the formatting could be better",
            "• some sentences are unclear and need improvement",
            "",
            "This document will be processed by the AI agent to:",
            "1. Fix grammar and spelling errors",
            "2. Improve clarity and readability",
            "3. Enhance formatting and structure", 
            "4. Maintain the original meaning",
            "",
            "After processing, the improved document will be saved and can be printed automatically.",
            "",
            "Additional test content:",
            "- dont forget to check this",
            "- cant wait to see the results",
            "- wont take long to process",
            "",
            "This sample demonstrates the AI agent's ability to improve documents while preserving their original intent and structure."
        ]
        
        for line in content:
            if line:
                doc.add_paragraph(line)
            else:
                doc.add_paragraph()
        
        # Save the document
        sample_path = Path("data/incoming/sample_document.docx")
        sample_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(sample_path)
        
        print(f"✅ Sample document created: {sample_path}")
        return sample_path
    
    if __name__ == "__main__":
        create_sample_document()
        
except ImportError:
    print("❌ python-docx not installed. Run: pip install python-docx")
except Exception as e:
    print(f"❌ Error creating sample document: {e}")