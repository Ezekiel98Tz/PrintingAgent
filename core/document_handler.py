"""
Document Handler for parsing, editing, and converting documents
Supports multiple formats: PDF, DOCX, DOC, TXT, RTF
"""

import logging
import mimetypes
from pathlib import Path
from typing import Optional, Dict, Any, List, Union
import tempfile
import shutil

# Document processing libraries
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    DOCX_AVAILABLE = False

try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    docx2txt = None
    DOCX2TXT_AVAILABLE = False

try:
    import PyPDF2
    from pypdf import PdfReader, PdfWriter
except ImportError:
    PyPDF2 = None
    PdfReader = None

try:
    import filetype
except ImportError:
    filetype = None

from config import Config

logger = logging.getLogger(__name__)

class DocumentHandler:
    """Handles document parsing, editing, and format conversion"""
    
    def __init__(self, config: Config):
        """Initialize document handler with configuration"""
        self.config = config
        self.supported_formats = config.supported_formats
        self.max_file_size = config.max_file_size_mb * 1024 * 1024  # Convert to bytes
        
    def validate_document(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Validate if document can be processed
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary with validation results
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {"valid": False, "error": "File does not exist"}
        
        # Check file size
        file_size = file_path.stat().st_size
        if file_size > self.max_file_size:
            return {
                "valid": False, 
                "error": f"File too large: {file_size / 1024 / 1024:.1f}MB (max: {self.config.max_file_size_mb}MB)"
            }
        
        # Check file format
        file_format = self._detect_file_format(file_path)
        if file_format not in self.supported_formats:
            return {
                "valid": False,
                "error": f"Unsupported format: {file_format}. Supported: {', '.join(self.supported_formats)}"
            }
        
        return {
            "valid": True,
            "format": file_format,
            "size_mb": file_size / 1024 / 1024,
            "filename": file_path.name
        }
    
    def _detect_file_format(self, file_path: Path) -> str:
        """Detect file format using multiple methods"""
        # Try using filetype library first
        if filetype:
            try:
                kind = filetype.guess(str(file_path))
                if kind:
                    return kind.extension.lower()
            except Exception:
                pass
        
        # Fallback to file extension
        extension = file_path.suffix.lower().lstrip('.')
        
        # Map common extensions
        extension_map = {
            'docx': 'docx',
            'doc': 'doc',
            'pdf': 'pdf',
            'txt': 'txt',
            'rtf': 'rtf'
        }
        
        return extension_map.get(extension, extension)
    
    def extract_text(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Extract text content from document
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        file_path = Path(file_path)
        
        # Validate document first
        validation = self.validate_document(file_path)
        if not validation["valid"]:
            return {"success": False, "error": validation["error"]}
        
        file_format = validation["format"]
        
        try:
            if file_format == "pdf":
                return self._extract_pdf_text(file_path)
            elif file_format in ["docx", "doc"]:
                return self._extract_docx_text(file_path)
            elif file_format == "txt":
                return self._extract_txt_text(file_path)
            elif file_format == "rtf":
                return self._extract_rtf_text(file_path)
            else:
                return {"success": False, "error": f"No handler for format: {file_format}"}
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return {"success": False, "error": str(e)}
    
    def _extract_pdf_text(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF file"""
        if not PdfReader:
            return {"success": False, "error": "PDF processing library not available"}
        
        try:
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                
                return {
                    "success": True,
                    "text": text.strip(),
                    "format": "pdf",
                    "pages": len(reader.pages),
                    "metadata": reader.metadata if hasattr(reader, 'metadata') else {}
                }
        except Exception as e:
            return {"success": False, "error": f"PDF extraction failed: {e}"}
    
    def _extract_docx_text(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX/DOC file"""
        try:
            # Try using python-docx first
            if DOCX_AVAILABLE and Document:
                doc = Document(file_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                
                return {
                    "success": True,
                    "text": text,
                    "format": "docx",
                    "paragraphs": len(doc.paragraphs)
                }
            
            # Fallback to docx2txt
            elif DOCX2TXT_AVAILABLE and docx2txt:
                text = docx2txt.process(str(file_path))
                return {
                    "success": True,
                    "text": text,
                    "format": "docx"
                }
            else:
                return {"success": False, "error": "DOCX processing library not available"}
                
        except Exception as e:
            return {"success": False, "error": f"DOCX extraction failed: {e}"}
    
    def _extract_txt_text(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                        return {
                            "success": True,
                            "text": text,
                            "format": "txt",
                            "encoding": encoding
                        }
                except UnicodeDecodeError:
                    continue
            
            return {"success": False, "error": "Could not decode text file"}
            
        except Exception as e:
            return {"success": False, "error": f"TXT extraction failed: {e}"}
    
    def _extract_rtf_text(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from RTF file (basic implementation)"""
        try:
            # Basic RTF text extraction (strips RTF formatting)
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                
                # Simple RTF parsing - remove control words
                import re
                # Remove RTF control words
                text = re.sub(r'\\[a-z]+\d*\s?', '', content)
                # Remove braces
                text = re.sub(r'[{}]', '', text)
                # Clean up whitespace
                text = re.sub(r'\s+', ' ', text).strip()
                
                return {
                    "success": True,
                    "text": text,
                    "format": "rtf"
                }
                
        except Exception as e:
            return {"success": False, "error": f"RTF extraction failed: {e}"}
    
    def save_processed_document(
        self,
        content: str,
        original_filename: str,
        output_format: Optional[str] = None,
        save_dir: Optional[Union[str, Path]] = None,
    ) -> Dict[str, Any]:
        """
        Save processed content as a new document
        
        Args:
            content: Processed text content
            original_filename: Original filename for reference
            output_format: Desired output format (defaults to config)
            
        Returns:
            Dictionary with save results and file path
        """
        output_format = (output_format or self.config.output_format).lower()
        
        try:
            # Determine output filename and directory
            original_path = Path(original_filename)
            output_dir = Path(save_dir) if save_dir else self.config.processed_dir
            try:
                output_dir.mkdir(parents=True, exist_ok=True)
            except Exception:
                pass

            # If caller provided a desired filename with matching suffix, honor it
            suffix = original_path.suffix.lower().lstrip('.')
            if suffix == output_format:
                output_filename = original_path.name
            else:
                output_filename = f"{original_path.stem}_processed.{output_format}"
            output_path = output_dir / output_filename
            
            if output_format == "txt":
                return self._save_as_txt(content, output_path)
            elif output_format == "docx":
                # Preserve formatting when requested and original is DOCX
                if self.config.preserve_formatting and original_path.suffix.lower() == ".docx":
                    preserve_result = self._save_as_docx_preserve_formatting(original_path, content, output_path)
                    # Fallback if preservation fails
                    if not preserve_result.get("success"):
                        return self._save_as_docx(content, output_path)
                    return preserve_result
                else:
                    return self._save_as_docx(content, output_path)
            elif output_format == "pdf":
                # For PDF, save as DOCX first then convert if needed
                docx_path = output_path.with_suffix('.docx')
                if self.config.preserve_formatting and original_path.suffix.lower() == ".docx":
                    docx_result = self._save_as_docx_preserve_formatting(original_path, content, docx_path)
                    if not docx_result.get("success"):
                        docx_result = self._save_as_docx(content, docx_path)
                else:
                    docx_result = self._save_as_docx(content, docx_path)
                if docx_result["success"]:
                    return {"success": True, "file_path": str(docx_path), "format": "docx"}
                return docx_result
            else:
                return {"success": False, "error": f"Unsupported output format: {output_format}"}
                
        except Exception as e:
            logger.error(f"Error saving processed document: {e}")
            return {"success": False, "error": str(e)}
    
    def _save_as_txt(self, content: str, output_path: Path) -> Dict[str, Any]:
        """Save content as TXT file"""
        try:
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(content)
            
            return {
                "success": True,
                "file_path": str(output_path),
                "format": "txt",
                "size_bytes": output_path.stat().st_size
            }
        except Exception as e:
            return {"success": False, "error": f"Failed to save TXT: {e}"}
    
    def _save_as_docx(self, content: str, output_path: Path) -> Dict[str, Any]:
        """Save content as DOCX file"""
        if not Document:
            return {"success": False, "error": "DOCX library not available"}

        try:
            doc = Document()
            
            # Split content into paragraphs and add to document
            paragraphs = content.split('\n\n')
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text.strip())
            
            doc.save(str(output_path))
            
            return {
                "success": True,
                "file_path": str(output_path),
                "format": "docx",
                "size_bytes": output_path.stat().st_size
            }
        except Exception as e:
            return {"success": False, "error": f"Failed to save DOCX: {e}"}

    def _save_as_docx_preserve_formatting(self, original_path: Path, content: str, output_path: Path) -> Dict[str, Any]:
        """
        Save content into a DOCX while preserving the original document's paragraph styles
        and run-level formatting, by editing the original file structure in place.

        Best-effort strategy:
        - Map content lines to existing paragraphs sequentially.
        - Distribute characters across existing runs to keep fonts/sizes.
        - Carry forward the last run's style for overflow text.
        """
        if not Document:
            return {"success": False, "error": "DOCX library not available"}
        try:
            doc = Document(str(original_path))
            lines = content.splitlines()

            def apply_text_to_runs(paragraph, new_text: str):
                runs = list(paragraph.runs)
                if not runs:
                    paragraph.add_run(new_text)
                    return
                remaining = new_text
                for run in runs:
                    if not remaining:
                        run.text = ""
                        continue
                    orig_len = len(run.text)
                    chunk_len = min(len(remaining), orig_len if orig_len > 0 else 32)
                    run.text = remaining[:chunk_len]
                    remaining = remaining[chunk_len:]
                if remaining:
                    last = runs[-1]
                    new_run = paragraph.add_run(remaining)
                    try:
                        new_run.style = last.style
                        new_run.bold = last.bold
                        new_run.italic = last.italic
                        new_run.underline = last.underline
                        if last.font is not None:
                            new_run.font.name = last.font.name
                            new_run.font.size = last.font.size
                            if getattr(last.font.color, "rgb", None) is not None:
                                new_run.font.color.rgb = last.font.color.rgb
                    except Exception:
                        pass

            line_idx = 0
            for para in doc.paragraphs:
                if line_idx >= len(lines):
                    break
                apply_text_to_runs(para, lines[line_idx].rstrip("\r"))
                line_idx += 1

            # Do not append extra paragraphs to avoid layout changes
            doc.save(str(output_path))
            return {
                "success": True,
                "file_path": str(output_path),
                "format": "docx",
                "size_bytes": output_path.stat().st_size
            }
        except Exception as e:
            return {"success": False, "error": f"Failed to save DOCX (preserve formatting): {e}"}
    
    def get_document_info(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Get comprehensive information about a document"""
        file_path = Path(file_path)
        
        validation = self.validate_document(file_path)
        if not validation["valid"]:
            return validation
        
        extraction = self.extract_text(file_path)
        
        info = {
            "filename": file_path.name,
            "format": validation["format"],
            "size_mb": validation["size_mb"],
            "valid": True
        }
        
        if extraction["success"]:
            text = extraction["text"]
            info.update({
                "word_count": len(text.split()),
                "character_count": len(text),
                "line_count": len(text.split('\n')),
                "has_content": bool(text.strip())
            })
            
            # Add format-specific info
            if "pages" in extraction:
                info["pages"] = extraction["pages"]
            if "paragraphs" in extraction:
                info["paragraphs"] = extraction["paragraphs"]
        else:
            info["extraction_error"] = extraction["error"]
        
        return info